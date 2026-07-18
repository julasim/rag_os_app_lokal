"""
Datei-Parser.

Jeder Parser gibt ein `ParsedDocument` zurück — eine einheitliche Repräsentation
mit Seiten/Abschnitten und Volltext. Das macht das nachfolgende Chunking
Format-unabhängig.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

import tempfile

import fitz  # PyMuPDF
import magic
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from openpyxl import load_workbook

from logger import log


# ---------------------------------------------------------------------------
# Datenmodell
# ---------------------------------------------------------------------------
@dataclass
class ParsedPage:
    """Eine logische Seite (PDF) oder Abschnitt (DOCX/XLSX)."""
    number: int
    text: str
    title: str | None = None
    title_level: int | None = None   # Überschriften-Ebene (1=oberste) für section_path


# Nummerierte Überschrift, z.B. "6.2.3 Kostenschätzung" → Ebene 3.
_NUMBERING_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)(?:\s|\)|\.)")


def _numbering_level(title: str | None) -> int | None:
    """Ebene aus der Gliederungsnummer (Anzahl der Punkt-Segmente). Sonst None."""
    if not title:
        return None
    m = _NUMBERING_RE.match(title)
    if not m:
        return None
    return m.group(1).count(".") + 1


@dataclass
class ParsedDocument:
    file_name: str
    mime_type: str
    pages: list[ParsedPage] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------
def parse_file(path: Path) -> ParsedDocument:
    mime = magic.from_file(str(path), mime=True)
    suffix = path.suffix.lower()

    log.info("ingest.parse.start", file=path.name, mime=mime)

    if mime == "application/pdf" or suffix == ".pdf":
        doc = _parse_pdf(path)
    elif suffix == ".docx" or "wordprocessingml" in mime:
        doc = _parse_docx(path)
    elif suffix in (".xlsx", ".xlsm") or "spreadsheetml" in mime:
        doc = _parse_xlsx(path)
    elif suffix in (".md", ".markdown"):
        doc = _parse_text(path, mime="text/markdown")
    elif suffix in (".txt", ".log"):
        doc = _parse_text(path, mime="text/plain")
    elif suffix in (".html", ".htm") or "html" in mime:
        doc = _parse_html(path)
    else:
        raise ValueError(f"Unsupported file type: {mime} ({suffix})")

    doc.mime_type = mime
    log.info(
        "ingest.parse.done",
        file=path.name,
        pages=len(doc.pages),
        chars=len(doc.full_text),
    )
    return doc


# ---------------------------------------------------------------------------
# Einzelne Parser
# ---------------------------------------------------------------------------
def _parse_pdf(path: Path) -> ParsedDocument:
    doc = _parse_pdf_pymupdf(path)
    # Gescannte PDFs haben keinen Text — OCR als Fallback
    if not doc.full_text.strip():
        log.info("ingest.parse.pdf.ocr_fallback", file=path.name)
        doc = _parse_pdf_ocr(path)
    return doc


def _parse_pdf_pymupdf(path: Path) -> ParsedDocument:
    """
    PDF-Parser mit visueller Leserichtung und Tabellen-Erkennung.

    Verbessert gegenüber get_text("text"):
    - get_text("blocks", sort=True) respektiert die visuelle Leserichtung
      (von oben-links nach unten-rechts) statt der internen Codierungsreihenfolge.
    - find_tables() erkennt tabellarische Strukturen und rendert sie als
      "Spalte1 | Spalte2 | Spalte3" pro Zeile — verhindert den Tabellen-Dump
      als kontextlose Wortliste.
    """
    doc = ParsedDocument(file_name=path.name, mime_type="application/pdf")
    with fitz.open(path) as pdf:
        for i, page in enumerate(pdf, start=1):
            parts: list[str] = []
            table_rects: list[fitz.Rect] = []

            # 1. Tabellen explizit als Zeilen mit Trennzeichen extrahieren
            try:
                for table in page.find_tables():
                    table_rects.append(fitz.Rect(table.bbox))
                    rows = table.extract()
                    rendered_rows: list[str] = []
                    for row in rows:
                        cells = [str(c or "").strip() for c in row if c is not None]
                        line = " | ".join(c for c in cells if c)
                        if line:
                            rendered_rows.append(line)
                    if rendered_rows:
                        parts.append("\n".join(rendered_rows))
            except Exception:
                # find_tables() nicht verfügbar oder fehlerhaft → ignorieren,
                # Tabellen werden dann als normale Blöcke verarbeitet
                table_rects = []

            # 2. Textblöcke in visueller Leserichtung (sort=True)
            #    Blöcke die in einer bereits extrahierten Tabelle liegen → überspringen
            for block in page.get_text("blocks", sort=True):
                # block = (x0, y0, x1, y1, text, block_no, block_type)
                if block[6] != 0:  # kein Text-Block (z.B. Bild)
                    continue
                block_rect = fitz.Rect(block[:4])
                if any(block_rect.intersects(tr) for tr in table_rects):
                    continue  # bereits als Tabelle verarbeitet
                text = block[4].strip()
                if text:
                    parts.append(text)

            full_text = "\n\n".join(parts).strip()
            if not full_text:
                continue

            first_line = full_text.splitlines()[0].strip()
            title = first_line if first_line and len(first_line) < 120 else None
            doc.pages.append(
                ParsedPage(number=i, text=full_text, title=title,
                           title_level=_numbering_level(title))
            )
    return doc


def _parse_pdf_ocr(path: Path) -> ParsedDocument:
    """OCR-gestütztes Parsen für gescannte PDFs via ocrmypdf + PyMuPDF."""
    try:
        import ocrmypdf
    except ImportError:
        log.warning("ingest.parse.pdf.ocr_unavailable", hint="pip install ocrmypdf")
        return ParsedDocument(file_name=path.name, mime_type="application/pdf")

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        ocr_path = Path(tmp.name)
    try:
        ocrmypdf.ocr(
            path,
            ocr_path,
            language="deu+eng",
            deskew=True,
            progress_bar=False,
        )
        return _parse_pdf_pymupdf(ocr_path)
    except Exception as e:
        log.warning("ingest.parse.pdf.ocr_failed", file=path.name, error=str(e))
        return ParsedDocument(file_name=path.name, mime_type="application/pdf")
    finally:
        ocr_path.unlink(missing_ok=True)


def _parse_docx(path: Path) -> ParsedDocument:
    doc = ParsedDocument(file_name=path.name, mime_type="")
    word = DocxDocument(str(path))

    current_title: str | None = None
    current_level: int | None = None
    buffer: list[str] = []
    section_num = 1

    def flush() -> None:
        nonlocal section_num, buffer, current_title, current_level
        text = "\n".join(buffer).strip()
        if text:
            doc.pages.append(
                ParsedPage(number=section_num, text=text,
                           title=current_title, title_level=current_level)
            )
            section_num += 1
        buffer = []

    for para in word.paragraphs:
        style = (para.style.name or "").lower()
        if style.startswith("heading") and para.text.strip():
            flush()
            current_title = para.text.strip()
            # Ebene aus dem Style ("heading 2" → 2), sonst aus der Gliederungsnummer.
            m = re.search(r"heading\s+(\d+)", style)
            current_level = int(m.group(1)) if m else _numbering_level(current_title)
            buffer.append(f"# {current_title}")
        else:
            if para.text.strip():
                buffer.append(para.text)
    flush()
    return doc


def _parse_xlsx(path: Path) -> ParsedDocument:
    doc = ParsedDocument(file_name=path.name, mime_type="")
    wb = load_workbook(path, data_only=True, read_only=True)
    for i, sheet in enumerate(wb.worksheets, start=1):
        rows = []
        for row in sheet.iter_rows(values_only=True):
            line = "\t".join("" if c is None else str(c) for c in row).rstrip()
            if line:
                rows.append(line)
        if rows:
            doc.pages.append(
                ParsedPage(number=i, title=sheet.title, text="\n".join(rows))
            )
    wb.close()
    return doc


def _parse_text(path: Path, mime: str) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    return ParsedDocument(
        file_name=path.name,
        mime_type=mime,
        pages=[ParsedPage(number=1, text=text)],
    )


def _parse_html(path: Path) -> ParsedDocument:
    html = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")
    for s in soup(["script", "style"]):
        s.decompose()
    text = soup.get_text(separator="\n")
    title = soup.title.string.strip() if soup.title and soup.title.string else None
    return ParsedDocument(
        file_name=path.name,
        mime_type="text/html",
        pages=[ParsedPage(number=1, text=text, title=title)],
    )
